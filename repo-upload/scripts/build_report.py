#!/usr/bin/env python3
"""
build_report.py - render a T3T consultation report into the real .dotx template.

The model (running the skill) produces the judgement-heavy part as JSON; THIS
script does the deterministic rendering so the output is always the exact T3T
template (Sage City boilerplate, sign-off, footer/logo all preserved).

Usage:
    python build_report.py report.json -o OUTPUT.docx [--template PATH]

report.json schema:
{
  "meta": {
    "company_name": "MLN",
    "date": "2026.06.24",
    "contact_person": "Lerusha Naidoo",
    "consultant": "Corne Pistorius",
    "case_nr": "345510"
  },
  "scope": "Needs analysis.",
  "actions": [
    {"text": "IRP5s not submitted for period 2025 - 2026"},
    {"text": "Load 2 Companies", "sub": ["Permanent", "Casuals"]},
    {"text": "Employee take-on sheet provided", "image": "shots/001.png",
     "caption": "Client email: take-on sheet"}
  ],
  "todo": [
    {"text": "Lerusha will mark the casual staff in red and send the sheet"}
  ],
  "instructions": [
    {"text": "On the next run, switch ACB to the FNB file for permanent staff"}
  ],
  "hours": {"hours": "1.0", "travel": "", "total": "1.0"}
}

Each actions/todo/instructions item: {text, sub?[...], image?, caption?}.

Notes:
- "instructions" is OPTIONAL. Use it only when the client/consultant gave
  forward instructions or recommendations that are distinct from agreed
  To-Do commitments (e.g. "next time, do X" / "we recommend Y"). Omit the
  key entirely if there are none. It renders as a "Recommendations /
  Instructions" section after To Do (and after Sources Used if present).
- Images are normalised through Pillow before embedding, so frames produced
  by ffmpeg (which lack standard JFIF headers and are otherwise rejected by
  python-docx) embed cleanly. Frame extraction is a first-class evidence
  source, not just a fallback.
"""

import argparse
import json
import os
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

try:
    from PIL import Image
    _HAVE_PIL = True
except ImportError:
    _HAVE_PIL = False


# --------------------------------------------------------------------------
# .dotx -> .docx  (python-docx rejects the template content-type)
# --------------------------------------------------------------------------
def template_to_docx(src, dst):
    if src.lower().endswith(".docx"):
        shutil.copy(src, dst)
        return dst
    tmp = tempfile.mkdtemp()
    with zipfile.ZipFile(src) as z:
        z.extractall(tmp)
    ct = os.path.join(tmp, "[Content_Types].xml")
    s = open(ct, encoding="utf-8").read()
    s = s.replace("wordprocessingml.template.main+xml",
                  "wordprocessingml.document.main+xml")
    open(ct, "w", encoding="utf-8").write(s)
    if os.path.exists(dst):
        os.remove(dst)
    with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(tmp):
            for f in files:
                fp = os.path.join(root, f)
                z.write(fp, os.path.relpath(fp, tmp))
    shutil.rmtree(tmp, ignore_errors=True)
    return dst


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------
def set_cell_text(cell, text):
    cell.text = text or ""


# Twips of space after each bullet/body paragraph (1 pt = 20 twips).
# 120 twips = 6pt, a clear but not loose gap between points.
_BULLET_SPACE_AFTER = 120


def _set_space_after(paragraph, twips=_BULLET_SPACE_AFTER):
    """Add a little breathing room after a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing")
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after",
        str(twips))


def _normalise_image(src_path, tmp_dir):
    """Re-save an image through Pillow so python-docx will accept it.

    ffmpeg-produced JPEGs lack the JFIF APP0 header python-docx's parser
    expects and are rejected with UnrecognizedImageError even though they
    are valid images. Round-tripping through Pillow adds a proper header.
    Returns a path python-docx can embed (the re-saved copy, or the original
    if Pillow isn't available / the image is already fine).
    """
    if not _HAVE_PIL:
        return src_path
    try:
        img = Image.open(src_path)
        img = img.convert("RGB")
        base = os.path.splitext(os.path.basename(src_path))[0]
        out = os.path.join(tmp_dir, base + "_norm.jpg")
        img.save(out, format="JPEG", quality=92)
        return out
    except Exception:
        return src_path


def insert_after(paragraph, text="", style=None, bullet_prefix="", space_after=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            np.style = style
        except Exception:
            text = bullet_prefix + text   # style missing -> fake the bullet
    elif bullet_prefix:
        text = bullet_prefix + text
    if text:
        np.add_run(text)
    if space_after:
        _set_space_after(np)
    return np


def find_para(doc, exact):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def insert_items_after(anchor, items, base_dir, img_tmp_dir=None):
    cur = anchor
    for it in items:
        has_children = bool(it.get("sub")) or bool(it.get("image"))
        cur = insert_after(cur, it["text"], style="List Bullet",
                            bullet_prefix="\u2022  ", space_after=not has_children)
        subs = it.get("sub", [])
        for i, sub in enumerate(subs):
            last_sub = (i == len(subs) - 1) and not it.get("image")
            cur = insert_after(cur, sub, style="List Bullet 2",
                               bullet_prefix="     \u25e6  ", space_after=last_sub)
        if it.get("image"):
            img = it["image"]
            if not os.path.isabs(img):
                img = os.path.join(base_dir, img)
            if img_tmp_dir:
                img = _normalise_image(img, img_tmp_dir)
            cur = insert_after(cur, "")
            run = cur.add_run()
            try:
                run.add_picture(img, width=Inches(5.8))
            except Exception as e:
                cur.add_run(f"[screenshot missing: {it['image']} - {e}]")
            if it.get("caption"):
                cur = insert_after(cur, it["caption"], style="Caption",
                                   bullet_prefix="", space_after=True)
            else:
                _set_space_after(cur)
    return cur


def insert_items(doc, after_heading, items, base_dir, img_tmp_dir=None):
    anchor = find_para(doc, after_heading)
    if anchor is None:
        raise RuntimeError(f"Heading not found in template: {after_heading!r}")
    return insert_items_after(anchor, items, base_dir, img_tmp_dir)


def fill_header(doc, meta):
    keymap = {
        "company name": "company_name", "date": "date",
        "contact person": "contact_person", "consultant": "consultant",
        "case nr": "case_nr", "site code": "site_code",
    }
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip().lower() == "company name":
            for row in t.rows:
                k = row.cells[0].text.strip().lower()
                if k in keymap and meta.get(keymap[k]) is not None and len(row.cells) > 1:
                    set_cell_text(row.cells[1], str(meta[keymap[k]]))
            return


def fill_hours(doc, hours):
    if not hours:
        return
    for t in doc.tables:
        hdr = [c.text.strip().lower() for c in t.rows[0].cells]
        if "hours/units" in hdr:
            for row in t.rows:
                c0 = row.cells[0].text.strip().lower()
                if c0 == "as per above":
                    set_cell_text(row.cells[2], str(hours.get("hours", "")))
                    if hours.get("details"):
                        set_cell_text(row.cells[1], str(hours["details"]))
                elif "travel" in row.cells[1].text.strip().lower():
                    set_cell_text(row.cells[2], str(hours.get("travel", "")))
                elif "total hours consulted" in c0:
                    set_cell_text(row.cells[2], str(hours.get("total", hours.get("hours", ""))))
            return


# --------------------------------------------------------------------------
def build(report, template, output):
    base_dir = report.get("_base_dir", os.path.dirname(os.path.abspath(output)))
    template_to_docx(template, output)
    doc = Document(output)

    fill_header(doc, report.get("meta", {}))

    img_tmp_dir = tempfile.mkdtemp(prefix="t3t_report_img_")
    try:
        # Scope: single line right after the heading
        scope = report.get("scope", "").strip()
        if scope:
            insert_after(find_para(doc, "Scope of Consultation:"), scope,
                         style="List Bullet", bullet_prefix="", space_after=True)

        insert_items(doc, "Actions Performed", report.get("actions", []),
                     base_dir, img_tmp_dir)
        last = insert_items(doc, "To Do:", report.get("todo", []),
                            base_dir, img_tmp_dir)

        # "Sources Used" extends the base template: insert it as a new section
        # after To Do (before the sign-off) whenever sources are declared.
        sources = report.get("sources", [])
        if sources:
            heading = insert_after(last, "Sources Used:", style="Heading 1")
            last = insert_items_after(heading, sources, base_dir, img_tmp_dir)

        # "Recommendations / Instructions" - optional. Only when the client or
        # consultant gave forward instructions/recommendations distinct from the
        # agreed To-Do commitments.
        instructions = report.get("instructions", [])
        if instructions:
            heading = insert_after(last, "Recommendations / Instructions:",
                                   style="Heading 1")
            insert_items_after(heading, instructions, base_dir, img_tmp_dir)

        fill_hours(doc, report.get("hours", {}))

        doc.save(output)
    finally:
        shutil.rmtree(img_tmp_dir, ignore_errors=True)
    return output


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("report_json")
    ap.add_argument("-o", "--output", required=True)
    ap.add_argument("--template", default=os.path.join(
        os.path.dirname(__file__), "..", "assets", "T3T_Consultation_report_Template.dotx"))
    args = ap.parse_args()
    report = json.load(open(args.report_json, encoding="utf-8"))
    report.setdefault("_base_dir", os.path.dirname(os.path.abspath(args.report_json)))
    out = build(report, args.template, args.output)
    print("Wrote", out)


if __name__ == "__main__":
    main()
